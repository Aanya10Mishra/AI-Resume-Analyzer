"""
Resume Parser
Extracts structured data from PDF and DOCX resumes
"""
import PyPDF2
import docx
import spacy
import re
from typing import Dict, List, Optional
import logging

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load spaCy model
try:
    nlp = spacy.load('en_core_web_sm')
    logger.info("✅ spaCy model loaded successfully")
except Exception as e:
    logger.error(f"❌ Failed to load spaCy model: {e}")
    logger.info("Please run: python -m spacy download en_core_web_sm")
    nlp = None

class ResumeParser:
    """
    Comprehensive resume parser for PDF and DOCX files
    Extracts: contact info, skills, experience, education
    """
    
    # Comprehensive skills database (expand as needed)
    SKILLS_DATABASE = {
        'programming_languages': [
            'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'c', 
            'ruby', 'php', 'swift', 'kotlin', 'go', 'rust', 'scala', 'r',
            'matlab', 'perl', 'shell', 'bash', 'powershell', 'dart', 'julia'
        ],
        'web_technologies': [
            'html', 'css', 'react', 'angular', 'vue', 'vue.js', 'svelte',
            'node.js', 'express', 'django', 'flask', 'fastapi', 'spring',
            'spring boot', 'asp.net', '.net', 'laravel', 'rails', 'ruby on rails',
            'next.js', 'nuxt', 'gatsby', 'redux', 'webpack', 'babel', 'sass',
            'less', 'bootstrap', 'tailwind', 'material-ui', 'jquery'
        ],
        'mobile_development': [
            'android', 'ios', 'react native', 'flutter', 'swift', 'kotlin',
            'xamarin', 'ionic', 'cordova', 'phonegap'
        ],
        'databases': [
            'sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch',
            'cassandra', 'dynamodb', 'oracle', 'mariadb', 'sqlite', 'neo4j',
            'couchdb', 'firebase', 'realm'
        ],
        'cloud_devops': [
            'aws', 'azure', 'gcp', 'google cloud', 'docker', 'kubernetes',
            'jenkins', 'ci/cd', 'terraform', 'ansible', 'vagrant', 'chef',
            'puppet', 'gitlab ci', 'github actions', 'circleci', 'travis ci',
            'cloudformation', 'helm', 'prometheus', 'grafana', 'nagios'
        ],
        'ml_ai': [
            'machine learning', 'deep learning', 'tensorflow', 'pytorch',
            'keras', 'scikit-learn', 'opencv', 'nlp', 'natural language processing',
            'computer vision', 'neural networks', 'cnn', 'rnn', 'lstm',
            'transformers', 'bert', 'gpt', 'data science', 'artificial intelligence',
            'reinforcement learning', 'xgboost', 'lightgbm', 'catboost'
        ],
        'data_analytics': [
            'pandas', 'numpy', 'matplotlib', 'seaborn', 'plotly', 'tableau',
            'power bi', 'excel', 'data analysis', 'statistics', 'stata',
            'spss', 'sas', 'data visualization', 'big data', 'hadoop',
            'spark', 'hive', 'pig', 'kafka', 'airflow'
        ],
        'version_control': [
            'git', 'github', 'gitlab', 'bitbucket', 'svn', 'mercurial'
        ],
        'project_management': [
            'jira', 'confluence', 'trello', 'asana', 'monday.com',
            'agile', 'scrum', 'kanban', 'waterfall', 'lean', 'safe'
        ],
        'testing': [
            'junit', 'pytest', 'jest', 'mocha', 'selenium', 'cypress',
            'cucumber', 'testng', 'unittest', 'jasmine', 'karma'
        ],
        'soft_skills': [
            'communication', 'leadership', 'teamwork', 'problem solving',
            'critical thinking', 'time management', 'presentation',
            'collaboration', 'adaptability', 'creativity', 'analytical',
            'detail-oriented', 'self-motivated', 'organizational'
        ]
    }
    
    def __init__(self):
        """Initialize parser with spaCy model"""
        self.nlp = nlp
        self.all_skills = self._flatten_skills()
    
    def _flatten_skills(self) -> List[str]:
        """Flatten skills database into single list"""
        all_skills = []
        for category, skills in self.SKILLS_DATABASE.items():
            all_skills.extend(skills)
        return all_skills
    
    # ============ FILE EXTRACTION ============
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text as string
        """
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            
            logger.info(f"✅ Extracted text from PDF ({num_pages} pages)")
            return text
        
        except Exception as e:
            logger.error(f"❌ Error extracting PDF: {e}")
            raise Exception(f"Error reading PDF: {str(e)}")
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text as string
        """
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            logger.info(f"✅ Extracted text from DOCX ({len(doc.paragraphs)} paragraphs)")
            return text
        
        except Exception as e:
            logger.error(f"❌ Error extracting DOCX: {e}")
            raise Exception(f"Error reading DOCX: {str(e)}")
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text based on file extension
        
        Args:
            file_path: Path to resume file
            
        Returns:
            Extracted text
        """
        extension = file_path.split('.')[-1].lower()
        
        if extension == 'pdf':
            return self.extract_text_from_pdf(file_path)
        elif extension in ['docx', 'doc']:
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    # ============ INFORMATION EXTRACTION ============
    
    def extract_email(self, text: str) -> Optional[str]:
        """Extract email address from text"""
        pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(pattern, text)
        return emails[0] if emails else None
    
    def extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number from text"""
        # Multiple phone patterns
        patterns = [
            r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # International
            r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # US format
            r'\d{10}',  # 10 digits
        ]
        
        for pattern in patterns:
            phones = re.findall(pattern, text)
            if phones:
                return phones[0]
        
        return None
    
    def extract_name(self, text: str) -> Optional[str]:
        """
        Extract candidate name using NLP
        Usually the first PERSON entity or first line
        """
        if not self.nlp:
            # Fallback: return first non-empty line
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            return lines[0] if lines else None
        
        # Use spaCy NER
        doc = self.nlp(text[:1000])  # Check first 1000 chars
        
        # Find first PERSON entity
        for ent in doc.ents:
            if ent.label_ == 'PERSON':
                return ent.text
        
        # Fallback
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        return lines[0] if lines else None
    
    def extract_linkedin(self, text: str) -> Optional[str]:
        """Extract LinkedIn URL"""
        pattern = r'(https?://)?(www\.)?linkedin\.com/in/[\w-]+'
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            return ''.join(matches[0]) if isinstance(matches[0], tuple) else matches[0]
        return None
    
    def extract_github(self, text: str) -> Optional[str]:
        """Extract GitHub URL"""
        pattern = r'(https?://)?(www\.)?github\.com/[\w-]+'
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            return ''.join(matches[0]) if isinstance(matches[0], tuple) else matches[0]
        return None
    
    def extract_skills(self, text: str) -> List[str]:
        """
        Extract technical and soft skills from resume
        Uses comprehensive skills database
        """
        text_lower = text.lower()
        found_skills = set()
        
        # Look for each skill in the text
        for skill in self.all_skills:
            # Create regex pattern for whole word matching
            pattern = r'\b' + re.escape(skill.lower()) + r'\b'
            if re.search(pattern, text_lower):
                found_skills.add(skill.title())  # Capitalize properly
        
        logger.info(f"✅ Found {len(found_skills)} skills")
        return sorted(list(found_skills))
    
    def extract_education(self, text: str) -> List[Dict]:
        """
        Extract education information
        Looks for degree keywords and surrounding context
        """
        education = []
        keywords = [
            'bachelor', 'master', 'phd', 'doctorate', 'b.tech', 'm.tech',
            'b.e', 'm.e', 'b.sc', 'm.sc', 'bca', 'mca', 'mba', 'degree',
            'diploma', 'associate', 'undergraduate', 'graduate', 'college',
            'university', 'institute', 'school'
        ]
        
        lines = text.split('\n')
        for i, line in enumerate(lines):
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in keywords):
                # Get surrounding context (previous and next 2 lines)
                context_start = max(0, i - 1)
                context_end = min(len(lines), i + 3)
                context = ' '.join(lines[context_start:context_end])
                
                education.append({
                    'degree': line.strip(),
                    'context': context.strip()
                })
        
        logger.info(f"✅ Found {len(education)} education entries")
        return education
    
    def extract_experience(self, text: str) -> List[Dict]:
        """
        Extract work experience
        Looks for job-related keywords and date patterns
        """
        experience = []
        keywords = [
            'experience', 'worked', 'working', 'intern', 'internship',
            'project', 'developed', 'built', 'designed', 'managed',
            'led', 'created', 'implemented', 'position', 'role',
            'engineer', 'developer', 'analyst', 'manager', 'coordinator'
        ]
        
        # Date patterns
        date_pattern = r'(19|20)\d{2}\s*[-–—to]\s*((19|20)\d{2}|present|current|ongoing)'
        
        lines = text.split('\n')
        for i, line in enumerate(lines):
            line_lower = line.lower()
            
            # Check for keywords or date patterns
            has_keyword = any(keyword in line_lower for keyword in keywords)
            has_date = re.search(date_pattern, line, re.IGNORECASE)
            
            if has_keyword or has_date:
                # Get surrounding context
                context_start = max(0, i - 1)
                context_end = min(len(lines), i + 4)
                context = ' '.join(lines[context_start:context_end])
                
                experience.append({
                    'title': line.strip(),
                    'context': context.strip()
                })
        
        logger.info(f"✅ Found {len(experience)} experience entries")
        return experience
    
    # ============ MAIN PARSING FUNCTION ============
    
    def parse(self, file_path: str) -> Dict:
        """
        Main parsing function
        Extracts all information from resume
        
        Args:
            file_path: Path to resume file
            
        Returns:
            Dictionary with all parsed information
        """
        try:
            logger.info(f"🔄 Starting to parse: {file_path}")
            
            # Extract raw text
            raw_text = self.extract_text(file_path)
            
            if not raw_text or len(raw_text.strip()) < 50:
                raise ValueError("Resume appears to be empty or too short")
            
            # Extract all components
            parsed_data = {
                'name': self.extract_name(raw_text),
                'email': self.extract_email(raw_text),
                'phone': self.extract_phone(raw_text),
                'linkedin': self.extract_linkedin(raw_text),
                'github': self.extract_github(raw_text),
                'skills': self.extract_skills(raw_text),
                'education': self.extract_education(raw_text),
                'experience': self.extract_experience(raw_text),
                'raw_text': raw_text,
                'word_count': len(raw_text.split()),
                'character_count': len(raw_text)
            }
            
            logger.info("✅ Resume parsed successfully!")
            logger.info(f"   - Name: {parsed_data['name']}")
            logger.info(f"   - Email: {parsed_data['email']}")
            logger.info(f"   - Skills: {len(parsed_data['skills'])} found")
            logger.info(f"   - Experience: {len(parsed_data['experience'])} entries")
            logger.info(f"   - Education: {len(parsed_data['education'])} entries")
            
            return parsed_data
        
        except Exception as e:
            logger.error(f"❌ Error parsing resume: {e}")
            raise Exception(f"Error parsing resume: {str(e)}")